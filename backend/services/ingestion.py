import os
import uuid
import logging
from datetime import datetime
from typing import Optional, Tuple, List
from pathlib import Path

import pypdf
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

from ..models import (
    FileType, DocumentStatus, DocumentMetadata, 
    DocumentChunk
)
from ..core.config import get_settings
from ..interfaces.services import IIngestionService
from ..interfaces.documents import IDocumentRepository

logger = logging.getLogger(__name__)

class IngestionService(IIngestionService):
    """Service for document ingestion and text extraction."""
    
    def __init__(self, document_repository: IDocumentRepository = None): 
        # Note: Repositories will be injected
        self.settings = get_settings()
        self.upload_dir = Path(self.settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.document_repository = document_repository

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, str]:
        """Validate file type and size."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in self.settings.allowed_extensions_list:
            return False, f"Invalid file type. Allowed: {', '.join(self.settings.allowed_extensions_list)}"
        
        if file_size > self.settings.max_file_size_bytes:
            return False, f"File too large. Maximum: {self.settings.max_file_size_mb}MB"
        
        return True, "Valid"
    
    async def save_file(self, filename: str, content: bytes) -> Tuple[str, str]:
        """Save uploaded file and return (unique_filename, file_path)."""
        ext = filename.rsplit(".", 1)[-1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{ext}"
        file_path = self.upload_dir / unique_filename
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        return unique_filename, str(file_path)
    
    async def extract_text(self, document_id: str) -> bool:
        """Extract text from document based on file type."""
        # Using the repository to find the document
        if not self.document_repository:
            logger.error("Document repository not initialized")
            return False

        doc = await self.document_repository.get(document_id)
        if not doc:
            logger.error(f"Document not found: {document_id}")
            return False
        
        try:
            # Update status to processing
            await self.document_repository.update_status(document_id, DocumentStatus.PROCESSING)
            
            # Since repository returns dict or model, handle accordingly.
            # Assuming dict for compatibility with existing logic for now, or adapt.
            # The repository implementation I wrote returns a dict (implicitly from motor)
            
            file_type_val = doc.get("file_type")
            file_path = doc.get("file_path")
            
            file_type = FileType(file_type_val)
            
            # Extract based on file type
            if file_type == FileType.PDF:
                text, metadata = self._extract_pdf(file_path)
            elif file_type == FileType.DOCX:
                text, metadata = self._extract_docx(file_path)
            elif file_type == FileType.XLSX:
                text, metadata = self._extract_xlsx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Create chunks for RAG
            chunks = self._create_chunks(text)
            
            # Update document with extracted content
            await self.document_repository.update_content(
                document_id, 
                text, 
                [c.model_dump() for c in chunks],
                metadata.model_dump() if metadata else None
            )
            
            logger.info(f"Successfully extracted text from document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to extract text from {document_id}: {e}")
            await self.document_repository.update_status(document_id, DocumentStatus.FAILED, str(e))
            return False
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from PDF file."""
        text_parts = []
        page_count = 0
        
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                page_count = len(reader.pages)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
        
        full_text =("\n".join(text_parts)).strip()
        
        # Fallback to pdfplumber
        if len(full_text) < 100:
             logger.info("Trying pdfplumber...")
             text_parts = []
             try:
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        text_parts.append(page.extract_text() or "")
                full_text = "\n".join(text_parts)
             except Exception as e:
                 logger.warning(f"pdfplumber failed: {e}")

        metadata = DocumentMetadata(pages=page_count, word_count=len(full_text.split()))
        return full_text, metadata
    
    def _extract_docx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text: paragraphs.append(row_text)
        
        full_text = "\n".join(paragraphs)
        return full_text, DocumentMetadata(word_count=len(full_text.split()))
    
    def _extract_xlsx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        workbook = load_workbook(file_path, data_only=True)
        text_parts = []
        sheet_names = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_names.append(sheet_name)
            text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) for cell in row if cell is not None]
                if any(v.strip() for v in row_values):
                    text_parts.append(" | ".join(row_values))
                    
        full_text = "\n".join(text_parts)
        return full_text, DocumentMetadata(sheets=sheet_names, word_count=len(full_text.split()))
    
    def _create_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[DocumentChunk]:
        if not text: return []
        chunks = []
        words = text.split()
        i = 0
        chunk_id = 0
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunks.append(DocumentChunk(chunk_id=chunk_id, text=" ".join(chunk_words)))
            chunk_id += 1
            i += chunk_size - overlap
        return chunks
